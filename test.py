from docx import Document
from io import BytesIO

def fill_word(data):
    doc=Document('template.docx')
    if doc.tables!=[]:
        for table in doc.tables:
            for cell in table._cells:
                for paragraph in cell.paragraphs:
                    tmp=paragraph.text
                    for i in range(len(data)):
                        comp='var'+str(i+1)
                        if comp in tmp.lower():
                            tmp=tmp.replace(comp,data[i])
                            paragraph.text=tmp
    return doc


data=['ijb', 'bi', 'hb', 'hb', '01.01.2001', '9:00', '12', 'bh', 'jhjh', 'b', 'jb', '2', '1', 'jhbhj', '1', 'bb', 'hjb', 'jh']
doc=fill_word(data)