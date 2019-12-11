import json
import os
import random
import smtplib
import datetime
from io import BytesIO
from datetime import date

from django.views.generic import ListView
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.core import serializers

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase

from ase_site.auth_core.models import User
from ase_site.data.models import Application, GPSdata, Car
from ase_site.data.choises import STATUS, TYPE
from .forms import beton_form, sand_pgs_form


class ViewAllRequests(ListView):
    template_name = "ase_site/req/templates/posts.html"
    model = Application

    def get_queryset(self):
        qs = super(ViewAllRequests, self).get_queryset()
        return qs


def approve2(request, post_id):
    document = Document()
    document.add_picture('root_files/ase_logo.png', width=Inches(1.25))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Дмитровское ш.,2,Москва', 3,)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Запрос!', level=1)
    req = Application.objects.get(id=post_id)
    document.add_paragraph(req.post)
    if req.status == 0:
        req.title = req.title+"(обработан)"
        req.status = 1
        req.save()
    else:
        req.save()
    document.save('root_files/test'+post_id+'.docx')
    #fromaddr = "testase@mail.ru"
    fromaddr = "test_mail_temp2018@mail.ru"
    toaddr = "dr.interned@gmail.com"
    msg = MIMEMultipart()
    msg['From'] = fromaddr
    msg['To'] = toaddr
    msg['Subject'] = "ЗАПРОС"
    body = "Вам пришел запрос"
    msg.attach(MIMEText(body, 'plain'))
    filename = "test.docx"
    attachment = open('root_files/test'+post_id+'.docx', "rb")
    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition',
                    "attachment; filename= %s" % filename)
    msg.attach(part)
    server = smtplib.SMTP('smtp.mail.ru', 587)
    server.starttls()
    #server.login(fromaddr, "adminase2018")
    server.login(fromaddr, "temp_pass2018")
    text = msg.as_string()
    server.sendmail(fromaddr, toaddr, text)
    server.quit()
    attachment.close()
    # time.sleep(15)
    os.remove("root_files/test"+post_id+".docx")
    return render(request, 'static/static_files/html/box.html', {'values': ['Запрос отправлен']})


def disapprove(request, post_id):
    req = Application.objects.get(id=post_id)
    req.current_level = req.current_level-1
    req.title = "Запрос №"+str(req.id)+"(отклонен)"
    req.save()
    return render(request, 'static/static_files/html/box.html', {'values': ['Запрос отклонены']})


def fill_word(data, doc_type):
    document = Document(os.path.join('docs_templates', doc_type+'_template.docx'))
    paragraphs = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)

    for p in paragraphs:
        if p.text in {str(i) for i in range(len(paragraphs))}:
            p.text = data[int(p.text)]
    return document


def create_request(request, application_type):
    if request.method == "POST":
        if application_type == 1:
            form = beton_form(request.POST)
        else: 
            form = sand_pgs_form(request.POST)
        if form.is_valid():
            application, _ = Application.objects.update_or_create(
                application_type=application_type,
                status=request.user.level * 2,
                density=form.cleaned_data['density'],
                delivery_place=request.POST.get('delivery_place'),
                delivery_date=request.POST.get('delivery_date'),
                delivery_time=request.POST.get('delivery_time'),
                volume=form.cleaned_data['volume'],

                project_number=request.POST.get('project_number'),
                material_class=request.POST.get('material_class'),
                antifreeze=request.POST.get('antifreeze'),
                water_resist=request.POST.get('water_resist'),
                freeze_resist=request.POST.get('freeze_resist'),
                ok=request.POST.get('ok'),
                lay_type=request.POST.get('lay_type'),
                delivery_type=request.POST.get('delivery_type'),
                
                car=form.cleaned_data['car'],
                manufacturer_org=form.cleaned_data['manufacturer_org'],
                performing_org=form.cleaned_data['performing_org'],
                application_sender=request.user,
                ocr_specialist=form.cleaned_data['ocr_specialist']
            )
            application.save()
            return redirect('/')
    else:
        if application_type == 1:
            form = beton_form()
        else: 
            form = sand_pgs_form()
        tp = {1: 'Бетон', 2: 'Песок', 3: 'ПГС'}
    return render(request, "ase_site/req/templates/index.html", {"form": form, "name": tp.get(application_type)})


def create_tuples(request, id_):
    application = Application.objects.get(id=id_)
    fields = application._meta.get_fields()
    data = []
    for f in fields:
        if not any(str(f.name) in s for s in ['status' 'id']):
            if str(f.name) == 'application_type':
                data.append(
                    (str(f.verbose_name), application.get_application_type_display()))
            else:
                if getattr(application, f.name) is not None:
                    try:
                        data.append((str(f.verbose_name), str(
                            getattr(application, f.name)), str(f.name)))
                    except:
                        pass
                elif f.name == 'application_receiver':
                    data.append((str(f.verbose_name), "", str(f.name)))
    return data


def get_coord_list(id_):
    try:
        connected_car = Car.objects.filter(connected_application=id_).first()
        coord_list = GPSdata.objects.filter(id_gps=connected_car.gps)
        final_coord = [val.get_tuple() for val in coord_list]
    except:
        final_coord = []
    return final_coord


def create_word(request, id_):
    application = Application.objects.get(id=id_)
    fields = application._meta.get_fields()
    data = []
    doc_type = {1: "beton", 2:"sand", 3:'pgs'}
    for f in fields:
        if not str(f.name) in ['status', 'application_type', 'id', 'connected_application']:
            if getattr(application, f.name) is not None:
                data.append(str(getattr(application, f.name)))
            elif f.name == 'application_receiver':
                data.append("")
            
    doc = fill_word(data, doc_type.get(application.application_type))
    data = BytesIO()
    doc.save(data)
    response = HttpResponse(data.getvalue(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="reports.docx"'
    return response


def create_beton_request(request):
    return create_request(request, 1)


def create_sand_request(request):
    return create_request(request, 2)


def create_PGS_request(request):
    return create_request(request, 3)


def show_application(request, id_, modal=False):
    app = Application.objects.get(id=id_)
    if request.method == "POST":
        comment = request.POST.get('comment')
        app.disapprove_comment = comment
        app.status = app.status - 1
        app.save()
        return redirect('/')
    else:
        full_application = create_tuples(request, id_)
        button_list = []
        user = request.user
        if user == app.application_sender and app.status == 8:
            button_list = ['complete', 'mistake']
        if user.level == 2 and app.status == 2:
            button_list = ['approve', 'disapprove']
        if user.level == 3 and app.status == 4:
            button_list = ['approve', 'disapprove']
        if user.level == 4 and app.status == 6:
            button_list = ['take_to_work']
        coord_list = get_coord_list(id_)
        return render(request, "ase_site/req/templates/post.html", 
        {'user': user,
         "application": full_application, 
         'app': app, 
         "coord_list": coord_list,
         'buttons': button_list,
         'modal': modal})


def show_all_applications(request, material_filter='all', status_filter='all', sort_field="id", sort_type="asc", mp=None):
    if request.method == "POST":
        return redirect('/request/all/%s/%s'%(request.POST.get('material_filter'), request.POST.get('status_filter')))
    sort_list = ['id', 'application_type', 'delivery_date', 'delivery_time', 'delivery_place', 'performing_org', 'status', 'compile_date', 'compile_time', 'application_sender']
    user = request.user
    if user.level == 1:
        applications = Application.objects.filter(application_sender=user).order_by(sort_field if sort_type == 'asc' else '-'+sort_field)
    elif user.level == 2:
        applications = Application.objects.filter(performing_org=user.firm_name).order_by(sort_field if sort_type == 'asc' else '-'+sort_field)
    elif user.level == 3:
        applications = Application.objects.filter(status__gte=3).order_by(sort_field if sort_type == 'asc' else '-'+sort_field)
    elif user.level == 4:
        applications = Application.objects.filter(status__gte=6).filter(manufacturer_org=user.firm_name).order_by(sort_field if sort_type == 'asc' else '-'+sort_field)
    else:
        applications = Application.objects.all().order_by(sort_field if sort_type == 'asc' else '-'+sort_field)
    if material_filter != 'all':
        applications = applications.filter(application_type=material_filter)
    if status_filter != 'all':
        applications = applications.filter(status=status_filter)
    if mp:
        list_of_coord_list = []
        for aplic in applications:
            coord_list = get_coord_list(aplic.id)
            r = lambda: random.randint(0,255)
            list_of_coord_list.append([coord_list, '#%02X%02X%02X' % (r(),r(),r())])
        return render(request, "ase_site/req/templates/posts_map.html",
        {'object_list': [(app, color[1]) for app, color in zip(applications, list_of_coord_list)],
        'list_of_coord_list': list_of_coord_list})
    return render(request, "ase_site/req/templates/posts.html", 
    {'object_list': applications, 
    "sort_type": ["asc" if f != sort_field or (f == sort_field and sort_type == 'desc') else "desc" for f in sort_list],
    "sort_option": sort_list,
    "sort_field": sort_field,
    "sort_field_type": sort_type,
    "material_filter":material_filter,
    "status_filter":status_filter,
    "material_list": TYPE,
    "status_list": STATUS})



def approve(request, id_):
    application = Application.objects.get(id=id_)
    application.status += 2
    if application.status == 8:
        available_cars = Car.objects.all().filter(car_type = application.car).filter(connected_application=None)
        try:
            connected_car = available_cars[random.randint(0, len(available_cars) - 1)]
            connected_car.connected_application = application
            connected_car.save()
            application.application_receiver = request.user
        except:
            application.status -= 2
            return show_application(request, id_=id_, modal=True)
    if application.status == 10:
        application.compile_date = date.today()
        application.compile_time = datetime.datetime.now().time().replace(microsecond=0)
        connected_car = Car.objects.filter(connected_application=id_).first()
        connected_car.connected_application = None
        connected_car.save()
    application.save()
    return redirect('/request/'+str(id_))
